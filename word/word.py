import os
import json
import time
import hashlib
import threading
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
import concurrent.futures
import sys
import string
from .word_style_utils import WordStyleCopier
from pdf2zh.translator import (
    OpenAITranslator,
    DeepLTranslator,
    GoogleTranslator,
    TencentTranslator,
    AzureTranslator,
    ZhipuTranslator,
    SiliconTranslator
)
from pdf2zh.cache import load_paragraph, write_paragraph, create_cache, deterministic_hash
from copy import deepcopy
from lxml import etree

# 并行翻译文本，同时保持原始顺序
def translate_text(index, text_item, translator, lang_from, lang_to, service):
    try:
        # 如果文本无效，直接返回原文
        if not check_text(text_item['text']):
            return {
                'index': index, 
                'text': text_item['text']
            }
        
        # 生成文本的哈希值
        text_hash = deterministic_hash(text_item['text'])
        
        # 尝试从缓存加载翻译
        cached_translation = load_paragraph(text_hash, f"{lang_from}_{lang_to}_{service}")

        if cached_translation:
            translated_text = cached_translation
        else:
            # 如果缓存未命中，进行翻译
            translated_text = translator.translate(text_item['text'])
            
            # 创建缓存目录
            create_cache(text_hash)
            
            # 保存翻译到缓存
            write_paragraph(text_hash, f"{lang_from}_{lang_to}_{service}", translated_text)
        
        return {
            'index': index, 
            'text': translated_text
        }
    except Exception as e:
        return {
            'index': index, 
            'text': text_item['text']  # 返回原文本
        }

def start(trans):
    """开始Word文档翻译"""
    
    # 允许的最大线程
    threads = trans.get('threads', 10)
    if threads is None or threads == "" or int(threads) < 0:
        max_threads = 10
    else:
        max_threads = int(threads)

    # 创建Document对象，加载Word文件
    try:
        file_path = trans.get('file_path', '')
        document = Document(file_path)
    except Exception as e:
        return False

    # 创建新文档并复制原始文档样式
    translated_document = Document()
    
    # 复制文档样式
    def deep_copy_document_styles(source_doc, target_doc):
        """深度复制文档样式，确保格式完全一致"""
        for style in source_doc.styles:
            try:
                if style.name not in target_doc.styles:
                    target_doc.styles.add_style(style.name, style.type)
            except Exception:
                pass

    deep_copy_document_styles(document, translated_document)

    texts = []
    trans_type = trans.get('type', 'trans_text_only_inherit')

    # 根据翻译类型读取文本
    if trans_type in ["trans_text_only_inherit", "trans_all_only_inherit", "trans_all_both_inherit"]:
        read_rune_text(document, texts)
    else:
        read_paragraph_text(document, texts)

    # 初始化翻译器
    service = trans.get('service', 'OpenAI')
    apikey = trans.get('apikey')
    model_id = trans.get('model_id', 'gpt-4o-mini')
    lang_from = trans.get('lang_from', 'Chinese')
    lang_to = trans.get('lang_to', 'English')

    # 选择翻译服务
    translator_map = {
        'OpenAI': OpenAITranslator,
        'DeepL': DeepLTranslator,
        'Google': GoogleTranslator,
        'Tencent': TencentTranslator,
        'Azure': AzureTranslator,
        'Zhipu': ZhipuTranslator,
        'Silicon': SiliconTranslator
    }
    
    translator_class = translator_map.get(service, OpenAITranslator)
    translator = translator_class(
        service=service,
        lang_out=lang_to, 
        lang_in=lang_from, 
        model=model_id
    )

    # 使用线程池进行并行翻译，并保持原始顺序
    translated_texts = [None] * len(texts)
    with concurrent.futures.ThreadPoolExecutor(max_workers=max_threads) as executor:
        # 提交所有文本片段到线程池，同时传递原始索引
        futures = [
            executor.submit(
                translate_text, 
                i, 
                text_item, 
                translator, 
                lang_from, 
                lang_to, 
                service
            ) 
            for i, text_item in enumerate(texts)
        ]
        
        # 按原始顺序收集结果
        for future in concurrent.futures.as_completed(futures):
            try:
                result = future.result()
                translated_texts[result['index']] = result
            except Exception as exc:
                pass

    # 替换原始文本列表为翻译后的文本
    texts = [item for item in translated_texts if item is not None]

    # 根据翻译类型写入文档
    text_count = 0
    write_type_map = {
        "trans_text_only_inherit": write_only_new,
        "trans_text_only_new": write_paragraph_text,
        "trans_text_both_new": write_both_new,
        "trans_text_both_inherit": write_rune_both,
        "trans_all_only_new": lambda doc, txts, cnt, only: write_paragraph_text(doc, txts, cnt, False),
        "trans_all_only_inherit": lambda doc, txts, cnt, only: write_only_new(doc, txts, cnt, False),
        "trans_all_both_new": lambda doc, txts, cnt, only: write_both_new(doc, txts, cnt, False),
        "trans_all_both_inherit": lambda doc, txts, cnt, only: write_rune_both(doc, txts, cnt, False)
    }

    write_func = write_type_map.get(trans_type)
    if write_func:
        # 修改传入的文本格式
        translated_texts_for_write = [{'text': t['text']} for t in texts]
        
        # 初始化文本计数器
        text_count = 0
        
        # 调用写入函数，并传递文本计数器
        text_count = write_func(translated_document, translated_texts_for_write, text_count, True)
        
        # 如果没有文本被写入，尝试强制写入
        if text_count == 0 and len(translated_texts_for_write) > 0:
            for paragraph in translated_document.paragraphs:
                paragraph.clear()
            
            for text in translated_texts_for_write:
                translated_document.add_paragraph(text['text'])

    # 复制媒体元素
    def copy_media_elements(source_doc, translated_document, trans):
        """
        复制媒体元素，保留翻译后的内容并添加图片和表格
        """
        
        # 保留已翻译的段落数量
        translated_paragraph_count = len(translated_document.paragraphs)
        
        # 复制图片
        for paragraph in source_doc.paragraphs:
            # 检查段落是否有图片
            has_image = any(
                run.element.find('.//pic:pic', namespaces={'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'}) is not None 
                for run in paragraph.runs
            )
            
            if has_image:
                # 找到对应的翻译段落（如果存在）
                matching_translated_paragraph = None
                paragraph_text = paragraph.text.strip()
                
                for trans_para in translated_document.paragraphs:
                    trans_para_text = trans_para.text.strip()
                    if trans_para_text and (trans_para_text in paragraph_text or paragraph_text in trans_para_text):
                        matching_translated_paragraph = trans_para
                        break
                
                # 如果找到匹配的段落，复制图片到该段落
                target_paragraph = matching_translated_paragraph or translated_document.paragraphs[-1]
                
                # 复制图片
                for run in paragraph.runs:
                    pic_element = run.element.find('.//pic:pic', namespaces={'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'})
                    if pic_element is not None:
                        WordStyleCopier.copy_image(target_paragraph, run)
        
        # 复制表格
        for source_table in source_doc.tables:
            # 使用精确复制方法，但不覆盖已翻译的内容
            copy_table_with_position(source_table, translated_document, trans)
        
        return translated_document

    def is_table_text(text_item):
        """
        判断文本是否属于表格内容的启发式方法
        """
        keywords = ['表', '序号', '项目', '日期', '金额', '编号', 'No.', 'Table']
        return any(keyword in text_item.get('original', '') for keyword in keywords)

    def copy_table_with_position(source_table, translated_document, trans):
        """
        精确复制表格，包括其在文档中的位置和格式，并插入翻译后的文本
        """
        
        # 专门处理表格内容的翻译映射
        table_translation_map = {}
        
        # 筛选表格相关的翻译文本
        table_translations = [
            text_item for text_item in trans 
            if is_table_text(text_item)
        ]
        
        # 创建翻译映射
        for text_item in table_translations:
            original = text_item.get('original', '').strip()
            translated = text_item.get('text', '').strip()
            if original:
                table_translation_map[original] = translated
        
        # 在翻译文档中创建新表格，复制源表格的样式和结构
        new_table = translated_document.add_table(
            rows=source_table.rows.__len__(), 
            cols=source_table.columns.__len__()
        )
        
        # 复制表格样式
        try:
            new_table.style = source_table.style
        except Exception:
            pass
        
        # 逐单元格复制和翻译
        for row_idx, source_row in enumerate(source_table.rows):
            for col_idx, source_cell in enumerate(source_row.cells):
                # 获取目标单元格
                target_cell = new_table.cell(row_idx, col_idx)
                
                # 清空目标单元格
                for para in target_cell.paragraphs:
                    para.clear()
                
                # 处理源单元格的段落
                for source_para in source_cell.paragraphs:
                    # 创建新段落
                    target_para = target_cell.add_paragraph()
                    
                    # 尝试复制段落样式
                    try:
                        target_para.style = source_para.style
                    except Exception:
                        pass
                    
                    # 处理每个run
                    for source_run in source_para.runs:
                        # 创建新run
                        target_run = target_para.add_run()
                        
                        # 复制run样式
                        try:
                            target_run.bold = source_run.bold
                            target_run.italic = source_run.italic
                            target_run.underline = source_run.underline
                            target_run.font.name = source_run.font.name
                            target_run.font.size = source_run.font.size
                        except Exception:
                            pass
                        
                        # 获取原始文本
                        original_text = source_run.text.strip()
                        
                        # 查找翻译（仅针对表格文本）
                        translated_text = table_translation_map.get(original_text, original_text)
                        
                        # 设置文本
                        target_run.text = translated_text
        
        return new_table

    # 复制媒体元素
    copy_media_elements(document, translated_document, texts)

    # 保存文档
    output_path = trans['file_path'].replace('.docx', '_translated.docx')
    translated_document.save(output_path)
    
    # 验证文档内容
    def verify_document_content(document_path, translated_texts):
        """验证生成文档的内容是否正确"""
        from docx import Document
        
        try:
            doc = Document(document_path)
            
            # 收集所有段落文本
            paragraph_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            
            # 检查是否所有翻译文本都在文档中
            missing_texts = [
                text['text'] for text in translated_texts 
                if not any(text['text'] in doc_text for doc_text in paragraph_texts)
            ]
            
            if missing_texts:
                return False
            
            return True
        
        except Exception as e:
            return False
    
    # 执行文档内容验证
    verify_result = verify_document_content(output_path, texts)
    if not verify_result:
        return False
    
    return True

def read_paragraph_text(document, texts):
    """按段落读取文本"""
    for paragraph in document.paragraphs:
        append_text(paragraph.text, texts)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                read_cell_text(cell, texts)

def read_rune_text(document, texts):
    """按run读取文本"""
    for paragraph in document.paragraphs:
        read_run(paragraph.runs, texts)
        
        # 处理超链接
        if paragraph.hyperlinks:
            for hyperlink in paragraph.hyperlinks:
                read_run(hyperlink.runs, texts)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                read_cell_text(cell, texts)

def read_cell_text(cell, texts):
    """读取单元格文本"""
    for paragraph in cell.paragraphs:
        append_text(paragraph.text, texts)

def read_run(runs, texts):
    """读取run文本"""
    for run in runs:
        append_text(run.text, texts)

def append_text(text, texts):
    """追加文本"""
    if check_text(text):
        texts.append({"text": text, "complete": False})

def check_text(text):
    """检查文本是否有效，包括处理空白和特殊字符"""
    if not text:
        return False
    
    # 去除所有空白字符
    stripped_text = text.strip()
    
    # 定义需要排除的特殊字符和文本
    special_chars = set('　_：:')
    exclude_texts = set()
    
    # 检查是否为空、仅包含特殊字符，或在排除列表中
    return (
        len(stripped_text) > 0 and 
        not all(char in special_chars for char in stripped_text) and
        stripped_text not in exclude_texts
    )

def write_paragraph_text(document, texts, text_count, onlyText):
    """按段落写入文本"""
    for paragraph in document.paragraphs:
        replace_paragraph_text(paragraph, texts, text_count, onlyText, False)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                write_paragraph_text(cell, texts, text_count, onlyText)

def write_both_new(document, texts, text_count, onlyText):
    """保留原文并重新排版"""
    for paragraph in document.paragraphs:
        replace_paragraph_text(paragraph, texts, text_count, onlyText, True)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                write_both_new(cell, texts, text_count, onlyText)

def write_only_new(document, texts, text_count, onlyText=True):
    """仅写入新文本，更加健壮的实现"""
    
    # 处理文档对象
    if isinstance(document, (DocxDocument, Document)):
        # 遍历文档段落
        for paragraph in document.paragraphs:
            # 清空段落
            paragraph.clear()
        
        # 重新填充段落
        for text_item in texts:
            current_text = text_item['text']
            document.add_paragraph(current_text)
        
        return len(texts)
    
    # 处理表格中的文本
    elif isinstance(document, (_Cell, Table)):
        for paragraph in document.paragraphs:
            paragraph.clear()
        
        for text_item in texts:
            current_text = text_item['text']
            document.add_paragraph(current_text)
        
        return len(texts)
    
    return text_count

def write_rune_both(document, texts, text_count, onlyText):
    """保留原文并继承版面"""
    for paragraph in document.paragraphs:
        if paragraph.runs:
            paragraph.runs[-1].add_break()
            add_paragraph_run(paragraph, paragraph.runs, texts, text_count)

        if paragraph.hyperlinks:
            for hyperlink in paragraph.hyperlinks:
                hyperlink.runs[-1].add_break()
                add_paragraph_run(paragraph, hyperlink.runs, texts, text_count)

        if onlyText:
            clear_image(paragraph)

        # 复制段落样式
        WordStyleCopier.copy_paragraph_style(paragraph, paragraph)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_paragraph_text(paragraph, texts, text_count, onlyText, True)

                    if paragraph.hyperlinks:
                        for hyperlink in paragraph.hyperlinks:
                            replace_paragraph_text(hyperlink, texts, text_count, onlyText, True)

def write_cell_text(cell, texts):
    """写入单元格文本"""
    for paragraph in cell.paragraphs:
        if check_text(paragraph.text) and texts:
            item = texts.pop(0)
            for index, run in enumerate(paragraph.runs):
                if index == 0:
                    run.text = item.get('text', "")
                else:
                    run.clear()

def write_run(runs, texts):
    """写入run文本"""
    text_count = 0
    if not runs:
        return text_count

    for run in runs:
        text = run.text
        if check_text(text) and texts:
            item = texts.pop(0)
            text_count += item.get('count', 0)
            run.text = item.get('text', "")

    return text_count

def add_paragraph_run(paragraph, runs, texts, text_count):
    """添加段落run"""
    for run in runs:
        if check_text(run.text) and texts:
            item = texts.pop(0)
            text_count += item.get('count', 0)
            new_run = paragraph.add_run(item.get('text', ""), run.style)
            set_run_style(new_run, run)
    
    set_paragraph_linespace(paragraph)

def set_run_style(new_run, copy_run):
    """设置run样式"""
    # 使用WordStyleCopier完整复制run样式
    WordStyleCopier.copy_run_style(new_run, copy_run)
    
    # 确保使用中文字体
    new_run.font.name = '微软雅黑'
    r = new_run._element.rPr.rFonts
    r.set(qn('w:eastAsia'), '微软雅黑')

def set_paragraph_linespace(paragraph):
    """设置段落行间距"""
    if hasattr(paragraph, "paragraph_format"):
        space_before = paragraph.paragraph_format.space_before
        space_after = paragraph.paragraph_format.space_after
        line_spacing = paragraph.paragraph_format.line_spacing
        line_spacing_rule = paragraph.paragraph_format.line_spacing_rule

        if space_before is not None:
            paragraph.paragraph_format.space_before = space_before
        if space_after is not None:
            paragraph.paragraph_format.space_after = space_after
        if line_spacing is not None:
            paragraph.paragraph_format.line_spacing = line_spacing
        if line_spacing_rule is not None:
            paragraph.paragraph_format.line_spacing_rule = line_spacing_rule

def check_image(run):
    """检查是否为图片"""
    return run.element.find('.//w:drawing', namespaces=run.element.nsmap) is not None

def clear_image(paragraph):
    """清除图片"""
    for run in paragraph.runs:
        if check_image(run):
            run.clear()

def replace_paragraph_text(paragraph, texts, text_count, onlyText, appendTo):
    """替换段落文本"""
    text = paragraph.text
    if check_text(text) and texts:
        item = texts.pop(0)
        trans_text = item.get('text', "")

        if appendTo:
            if paragraph.runs:
                paragraph.runs[-1].add_break()
                paragraph.runs[-1].add_text(trans_text)
            elif paragraph.hyperlinks:
                paragraph.hyperlinks[-1].runs[-1].add_break()
                paragraph.hyperlinks[-1].runs[-1].add_text(trans_text)
        else:
            replaced = False
            if paragraph.runs:
                for run in paragraph.runs:
                    if not check_image(run):
                        if not replaced:
                            run.text = trans_text
                            replaced = True
                        else:
                            run.clear()
            elif paragraph.hyperlinks:
                for hyperlink in paragraph.hyperlinks:
                    for run in hyperlink.runs:
                        if not check_image(run):
                            if not replaced:
                                run.text = trans_text
                                replaced = True
                            else:
                                run.clear()

        text_count += item.get('count', 0)
        set_paragraph_linespace(paragraph)

    if onlyText:
        clear_image(paragraph)

    return text_count
